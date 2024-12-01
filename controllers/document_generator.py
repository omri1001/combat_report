# document_generator.py

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

from app.models.chart_utils import create_bar_chart, create_final_grade_chart
from app.models.hyperlink_utils import add_hyperlink
from docx.enum.table import WD_TABLE_ALIGNMENT

def generate_word_document(sections, output_path, date="", signature="", title="", grades_data=None, youtube_link=None, poll_link=None):
    """
    Generates a Word document with the given content, applying the provided template.
    """
    # Load the template document
    template_path = '../views/tamplates/template.docx'  # Ensure this file exists in your directory
    if os.path.exists(template_path):
        document = Document(template_path)
    else:
        print(f"Error: Template file '{template_path}' not found. Using a blank document instead.")
        document = Document()

    # Set default font to support Hebrew characters
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'  # Ensure this font supports Hebrew
    font.size = Pt(10)
    font._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')  # Fix for Hebrew characters

    # --- First Page Content ---

    # Add the title centered at the top
    title_text = title if title else "דוח סיכום אימון"
    if title_text:
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align
        title_run = title_paragraph.add_run(title_text)
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(18)
        title_run.bold = True
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Add a line break before starting the GPT output text
    document.add_paragraph()

    # --- Insert GPT Output Text (Sections) ---
    # Add each section to the document
    for section_title, hebrew_title in [
        ("Introduction", "הקדמה"),
        ("Exercise 1", "תרגיל 1"),
        ("Exercise 2", "תרגיל 2"),
        ("Summary", "סיכום")
    ]:
        if sections.get(section_title):
            # Add section title
            heading = document.add_heading(hebrew_title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align
            heading_run = heading.runs[0]
            heading_run.font.name = 'Arial'
            heading_run.font.size = Pt(14)
            heading_run.bold = True
            heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

            # Add section content
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align
            content = sections[section_title]
            # Remove '#' and '*' from content
            content = content.replace('#', '').replace('*', '')
            run = paragraph.add_run(content)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # --- Grades Section ---
    if grades_data:
        # Add a page break before grades
        document.add_page_break()

        # Add Grades Section
        grades_heading = document.add_heading('דוח ציונים', level=1)
        grades_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        grades_heading_run = grades_heading.runs[0]
        grades_heading_run.font.name = 'Arial'
        grades_heading_run.font.size = Pt(14)
        grades_heading_run.bold = True
        grades_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # Insert graphs with comments between them
        max_items = max(len(part_data['items']) for part_name, part_data in grades_data.items() if part_name != 'final_grade')

        # Use an index to create filenames
        part_index = 1

        for part_name, part_data in grades_data.items():
            if part_name == "final_grade":
                continue
            # Add part heading
            part_heading = document.add_heading(part_name, level=2)
            part_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            part_heading_run = part_heading.runs[0]
            part_heading_run.font.name = 'Arial'
            part_heading_run.font.size = Pt(12)
            part_heading_run.bold = True
            part_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

            # Generate and insert graph for the part
            # Use index for the filename to avoid non-ASCII characters
            chart_path = f"part_{part_index}_chart.png"
            create_bar_chart(part_data['items'], part_name, chart_path, max_items)
            document.add_picture(chart_path, width=Inches(6))
            # Remove the chart image file after inserting
            os.remove(chart_path)

            part_index += 1  # Increment the index

            # Add comment after the graph
            if 'comment' in part_data and part_data['comment']:
                comment_paragraph = document.add_paragraph()
                comment_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align
                comment_run = comment_paragraph.add_run(part_data['comment'])
                comment_run.font.name = 'Arial'
                comment_run.font.size = Pt(10)
                comment_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # Add final grade
        final_grade_paragraph = document.add_paragraph()
        final_grade_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = final_grade_paragraph.add_run(f"ציון סופי: {grades_data['final_grade']}")
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        # Generate and insert final grade graph
        final_chart_path = "final_grade_chart.png"
        create_final_grade_chart(grades_data, final_chart_path, max_items)
        document.add_picture(final_chart_path, width=Inches(6))
        # Remove the chart image file after inserting
        os.remove(final_chart_path)

    # --- Final Page ---
    # Add a page break before the final page
    document.add_page_break()

    # Add final page content
    thank_you_title = document.add_heading("תודה שהשתתפתם באימון שלנו", level=0)
    thank_you_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    thank_you_title_run = thank_you_title.runs[0]
    thank_you_title_run.font.name = 'Arial'
    thank_you_title_run.font.size = Pt(24)
    thank_you_title_run.bold = True
    thank_you_title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Thank you message
    thank_you_message = (
        "אנו מודים לכם על השתתפותכם באימון שלנו. "
        "נשמח לעמוד לשירותכם בכל עת."
    )
    message_paragraph = document.add_paragraph()
    message_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    message_run = message_paragraph.add_run(thank_you_message)
    message_run.font.name = 'Arial'
    message_run.font.size = Pt(14)
    message_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Add the links after the thank you message with icons
    if youtube_link or poll_link:
        links_paragraph = document.add_paragraph()
        links_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if youtube_link:
            # Add YouTube icon
            youtube_icon = "../views/icons/icon_youtube.png"  # Ensure this file exists
            if os.path.exists(youtube_icon):
                youtube_run = links_paragraph.add_run()
                youtube_run.add_picture(youtube_icon, width=Inches(0.2))
            else:
                print(f"Warning: Icon file '{youtube_icon}' not found. Skipping icon.")
            # Add hyperlink
            add_hyperlink(links_paragraph, youtube_link, " סרטון התרגיל")
            links_paragraph.add_run("\n")

        if poll_link:
            # Add Poll icon
            poll_icon = "../views/icons/icon_poll.png"  # Ensure this file exists
            if os.path.exists(poll_icon):
                poll_run = links_paragraph.add_run()
                poll_run.add_picture(poll_icon, width=Inches(0.2))
            else:
                print(f"Warning: Icon file '{poll_icon}' not found. Skipping icon.")
            # Add hyperlink
            add_hyperlink(links_paragraph, poll_link, " סקר")
            links_paragraph.add_run("\n")

    # Contact information
    contact_info = [
        ("אתר האינטרנט", "https://digitalcombat.academy/"),
        ("דוא\"ל", "or.ben.shabat@digitalcombat.academy"),
        ("טלפון", "+972544538973"),
        ("LinkedIn", "https://www.linkedin.com/company/digital-combat-academy-incorporate/")
    ]

    # Add a table for contact information
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for label, info in contact_info:
        row_cells = table.add_row().cells
        # Icon cell
        icon_cell = row_cells[0]
        icon_paragraph = icon_cell.paragraphs[0]
        icon_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Map label to icon filename
        icon_filename = None
        if label == "אתר האינטרנט":
            icon_filename = "../views/icons/icon_website.png"
        elif label == "דוא\"ל":
            icon_filename = "../views/icons/icon_email.png"
        elif label == "טלפון":
            icon_filename = "../views/icons/icon_phone.png"
        elif label == "LinkedIn":
            icon_filename = "../views/icons/icon_linkedin.png"

        if icon_filename and os.path.exists(icon_filename):
            run = icon_paragraph.add_run()
            run.add_picture(icon_filename, width=Inches(0.2))
        else:
            print(f"Warning: Icon file '{icon_filename}' not found. Skipping icon.")

        # Text cell
        text_cell = row_cells[1]
        text_paragraph = text_cell.paragraphs[0]
        text_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if label != "טלפון":
            # Add clickable hyperlink
            add_hyperlink(text_paragraph, info, f"{label}: {info}")
        else:
            # Add regular text for the phone number
            contact_run = text_paragraph.add_run(f"{label}: {info}")
            contact_run.font.name = 'Arial'
            contact_run.font.size = Pt(12)
            contact_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Add the logo image if needed (assuming you still want to keep dca_logo2.png)
    logo2_path = "../views/icons/dca_logo2.png"  # Ensure this file exists in your directory
    if os.path.exists(logo2_path):
        logo_paragraph = document.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_run = logo_paragraph.add_run()
        logo_run.add_picture(logo2_path, width=Inches(6))
    else:
        print(f"Warning: Logo file '{logo2_path}' not found. Skipping logo.")

    # Add signature at the end
    if signature:
        document.add_paragraph()  # Add empty paragraph for spacing
        signature_paragraph = document.add_paragraph()
        signature_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Center align
        run = signature_paragraph.add_run(signature)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

    # Save the document
    document.save(output_path)
